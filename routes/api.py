from models import TableRow, DocumentData, AudioText, TimeAction, AudioAnalysisResponse
from config import UPLOAD_DIR
import io
import os
import shutil
import time
from datetime import datetime
from tempfile import NamedTemporaryFile
from typing import List

# FastAPI и Starlette
from fastapi import (
    FastAPI,
    APIRouter,
    File,
    Query,
    Request,
    UploadFile,
    WebSocket,
    status,
)

from fastapi.background import BackgroundTasks
from fastapi.exceptions import HTTPException
from fastapi.middleware.httpsredirect import HTTPSRedirectMiddleware
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.middleware.cors import CORSMiddleware
from starlette.websockets import WebSocketDisconnect, WebSocketState

# Сторонние библиотеки
import httpx
import requests
from docx import Document
from docx.shared import Inches
from fpdf import FPDF

# Логирование
import logging

SERVER_URL = "https://d776-193-41-143-66.ngrok-free.app"

router = APIRouter()


@router.get("/", response_class=HTMLResponse)
async def get():
    with open("templates/index.html", "r", encoding="utf-8") as f:
        return f.read()


@router.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    try:
        await websocket.accept()
        while True:
            try:
                data = await websocket.receive_text()
                await websocket.send_text(data)
            except WebSocketDisconnect:
                print("Client disconnected")
                break
            except Exception as e:
                print(f"Error in websocket communication: {e}")
                break
    finally:
        try:
            await websocket.close()
        except:
            pass

@router.post("/uploadfile/")
async def upload_file(file: UploadFile = File(...)):
    try:
        # Проверка типа файла
        if not file.content_type.startswith("audio/"):
            raise HTTPException(status_code=400, detail="Файл должен быть аудио формата.")

        # Сохраняем файл на сервере
        file_location = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_location, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Логируем путь файла
        print(f"Файл сохранен: {file_location}")

        # Путь к сохраненному файлу
        file_path = file_location

        external_api_url = f"{SERVER_URL}/transcribe/"
        async with httpx.AsyncClient(verify=False, timeout=30.0) as client:  # добавляем timeout
            with open(file_location, 'rb') as f:
                time.sleep(1)
                files = {'file': (file.filename, f, file.content_type)}
                response = await client.post(external_api_url, files=files)
                print(f"Тело ответа: {response.text}")

        # Проверяем ответ
        if response.status_code == 200:
            response_data = response.json()
            # Проверяем наличие текста в ответе
            if "text" in response_data:
                return {
                    "status": "success",
                    "transcribed_text": response_data["text"]
                }
            # Если есть ошибка в ответе
            elif "error" in response_data:
                return {
                    "status": "error",
                    "message": response_data["error"]
                }
            else:
                return {
                    "status": "error",
                    "message": "Неожиданный формат ответа от сервера"
                }
        else:
            return {
                "status": "error",
                "message": f"Ошибка сервера: {response.status_code}"
            }

    except Exception as e:
        return {
            "status": "error",
            "message": f"Произошла ошибка: {str(e)}"
        }
    finally:
        # Удаляем временный файл
        if file_location and os.path.exists(file_location):
            time.sleep(1)
            os.remove(file_location)

@router.post("/analyze_audio", response_model=AudioAnalysisResponse)
async def analyze_audio(audio_data: AudioText):
    try:
        external_api_url = f"{SERVER_URL}/analyze_audio"

        # Отправляем запрос к внешнему API
        async with httpx.AsyncClient() as client:
            response = await client.post(
                external_api_url,
                json=audio_data.dict()  # Преобразуем входные данные в JSON
            )

            # Проверяем статус ответа
            response.raise_for_status()

            # Получаем данные из ответа
            data = response.json()

            print(data)

            # Форматируем ответ в соответствии с моделью AudioAnalysisResponse
            formatted_events = [
                TimeAction(
                    start=event["start"],
                    end=event["end"],
                    action=event["action"]
                )
                for event in data["events"]
            ]
            print(formatted_events)

            return AudioAnalysisResponse(events=formatted_events).dict()

    except httpx.HTTPError as e:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Error communicating with external API: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Internal server error: {str(e)}"
        )


@router.post("/generate-docx")
async def generate_docx(data: DocumentData):
    try:
        # Создаем новый документ
        doc = Document()

        # Добавляем заголовок
        doc.add_heading("Распознавание текста получено при помощи сервиса DiaBert", level=1)

        # Подзаголовок: Распознанный текст
        doc.add_heading("Распознанный текст", level=2)

        # Проверяем, есть ли текст
        if data.text.strip():
            # Если текст есть, добавляем его
            doc.add_paragraph(data.text)
        else:
            # Если текста нет, выводим сообщение
            doc.add_paragraph("Текст не был распознан.")

        # Подзаголовок: Распознанный текст
        doc.add_heading("Таблица с расписанием операции и временем:", level=2)

        # Добавляем таблицу
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Заголовки таблицы
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Операция'
        header_cells[1].text = 'Время начала'
        header_cells[2].text = 'Время завершения'

        # Добавляем данные в таблицу
        for row_data in data.tableData:
            row_cells = table.add_row().cells
            row_cells[0].text = row_data.operation
            row_cells[1].text = row_data.startTime
            row_cells[2].text = row_data.endTime

        # Создаем временный файл
        with NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        # Возвращаем файл и удаляем его после отправки
        return FileResponse(
            tmp_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename='time_tracking.docx'
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))